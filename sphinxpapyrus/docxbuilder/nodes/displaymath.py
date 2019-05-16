# -*- coding: utf-8 -*-
"""
Translate docutils node displaymath formatting.
each displaymath start will processed with visit() and finished with depart()
"""

from docutils.nodes import Node
from docx.enum.table import WD_TABLE_ALIGNMENT

from sphinxpapyrus.docxbuilder.translator import DocxTranslator

node_name = "displaymath"


def visit(visitor: DocxTranslator, node: Node):
    """Start processing displaymath node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Node)

    eq = node.get('latex')
    eq = eq if eq else node.astext()
    number = node.get('number')
    number = '(%s)' % str(number) if number else ''
    visitor._add_paragraph_between_table(node)
    table = visitor.p_parents[-1].add_table(rows=1, cols=3)
    twidth = sum([cell.width for cell in table.row_cells(0)])
    table.cell(0, 0).width = int(twidth * 0.1)
    table.cell(0, 0).text = ''
    table.cell(0, 0).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
    table.cell(0, 1).width = int(twidth * 0.8)
    table.cell(0, 1).text = eq
    table.cell(0, 1).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 2).width = int(twidth * 0.1)
    table.cell(0, 2).text = number
    table.cell(0, 2).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT
    

def depart(visitor: DocxTranslator, node: Node):
    """Finish processing displaymath node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Node)

    visitor.p = None
