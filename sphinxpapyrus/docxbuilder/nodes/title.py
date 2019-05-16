# -*- coding: utf-8 -*-
"""
Translate docutils node title formatting.
each title start will processed with visit() and finished with depart()
"""
from docutils import nodes
from docutils.nodes import Node
from docx.enum.text import WD_BREAK

from sphinxpapyrus.docxbuilder.templates.styles.definitions import STYLES
from sphinxpapyrus.docxbuilder.templates.styles.names import TABLE_CAPTION
from sphinxpapyrus.docxbuilder.translator import DocxTranslator

node_name = "title"


def visit(visitor: DocxTranslator, node: Node):
    """Start processing title node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Node)

    if isinstance(node.parent, nodes.topic):
        pass

    elif isinstance(node.parent, nodes.sidebar):
        pass

    elif isinstance(node.parent, nodes.Admonition):
        pass

    elif isinstance(node.parent, nodes.table):
        prefix = visitor._fignum_prefix(node.parent)
        visitor.p = visitor._add_paragraph(
            prefix + ' ',
            style=STYLES[TABLE_CAPTION]
        )
        visitor.p.paragraph_format.keep_with_next = True

    elif isinstance(node.parent, nodes.document):
        visitor.docx.add_heading(node.astext().replace('\n', ' '), 0)

    elif isinstance(node.parent, nodes.section):
        headinglevel = visitor.section_level + visitor.initial_header_level - 1
        breaklevel = visitor.builder.config.docx_pagebreak_level
        if breaklevel is not None and headinglevel <= breaklevel:
            lastp = visitor.docx.paragraphs[-1] if len(
                visitor.docx.paragraphs) > 0 else None
            if lastp:
                lastp.add_run().add_break(WD_BREAK.PAGE)

        p = visitor.docx.add_heading(node.astext().replace('\n', ' '),
                                  headinglevel)
        secnumlevel = visitor.section_level - visitor.numbered_level
        if visitor.numbered and visitor.numbered > secnumlevel - 1:
            visitor._multilevel_list_numbering(p, secnumlevel - 1,
                                            visitor.section_numIds[-2])
    else:
        pass


def depart(visitor: DocxTranslator, node: Node):
    """Finish processing title node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Node)

    visitor.p = None
