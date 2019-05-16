# -*- coding: utf-8 -*-
"""
Translate docutils node tgroup formatting.
each tgroup start will processed with visit() and finished with depart()
"""
from docutils.nodes import Element
from docutils.nodes import Node
from docutils.nodes import tgroup
from docx.enum.table import WD_TABLE_ALIGNMENT

from sphinxpapyrus.docxbuilder.templates.styles.definitions import STYLES
from sphinxpapyrus.docxbuilder.templates.styles.names import TABLE_LIST
from sphinxpapyrus.docxbuilder.templates.styles.names import TABLE_NORMAL
from sphinxpapyrus.docxbuilder.translator import DocxTranslator

node_name = "tgroup"


def visit(visitor: DocxTranslator, node: Element):
    """Start processing tgroup node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Element)

    tgroup_node = node
    assert isinstance(tgroup_node, tgroup)

    # TODO: check if tgroup_node.children elements are Element
    thead_num = len(tgroup_node.children[-2].children)
    tbody_num = len(tgroup_node.children[-1].children)
    row_num = thead_num + tbody_num
    col_num = tgroup_node['cols']

    new_table = visitor.p_parents[-1].add_table(rows=row_num, cols=col_num)
    align = tgroup_node.parent.get('align')
    if not align:
        align = visitor.builder.config.docx_imagetable_align

    if align:
        if align == 'left':
            new_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        elif align == 'center':
            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif align == 'right':
            new_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    for raw in new_table.rows:
        for c in raw._tr.tc_lst:
            tcW = c.tcPr.tcW
            tcW.type = 'auto'
            tcW.w = 0

    if thead_num == 0:
        new_table.style = STYLES[TABLE_NORMAL]

    else:
        new_table.style = STYLES[TABLE_LIST]
        # for i in range(thead_num):
        #     trPr = table.rows[i]._tr.get_or_add_trPr()
        #     if isinstance(trPr, CT_TrPr):
        #         tblHeader = trPr.get_or_add_tblHeader()

    visitor.tables.append([new_table, 0, 0])


def depart(visitor: DocxTranslator, node: Node):
    """Finish processing tgroup node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Node)

    visitor.tables.pop()
