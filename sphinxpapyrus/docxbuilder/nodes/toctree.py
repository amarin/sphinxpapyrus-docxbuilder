# -*- coding: utf-8 -*-
"""
Translate docutils node toctree formatting.
each toctree start will processed with visit() and finished with depart()
"""

from docutils.nodes import Node
from docx.enum.text import WD_BREAK
from docx.opc.oxml import qn
from docx.oxml import OxmlElement

from sphinxpapyrus.docxbuilder.translator import DocxTranslator

node_name = "toctree"


def visit(visitor: DocxTranslator, node: Node):
    """Start processing toctree node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Node)

    # TODO: check docx.add_paragraph() method
    toc_levels = "1-3"
    paragraph = visitor.docx.add_paragraph()
    visitor.r = run = paragraph.add_run()

    # creates a new element
    char_begin = OxmlElement('w:fldChar')
    char_begin.set(qn('w:fldCharType'), 'begin')

    instrument = OxmlElement('w:instrText')
    instrument.set(qn('xml:space'), 'preserve')

    # change 1-3 depending on heading levels you need
    instrument.text = r'TOC \o "%s" \h \z \u' % toc_levels

    char_separate = OxmlElement('w:fldChar')
    char_separate.set(qn('w:fldCharType'), 'separate')

    field_char_3 = OxmlElement('w:t')
    field_char_3.text = "Right-click to update field."

    char_separate.append(field_char_3)

    char_end = OxmlElement('w:fldChar')
    char_end.set(qn('w:fldCharType'), 'end')

    run.element.append(char_begin)
    run.element.append(instrument)
    run.element.append(char_separate)
    run.element.append(char_end)
    

def depart(visitor: DocxTranslator, node: Node):
    """Finish processing toctree node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Node)

    visitor.r.add_run().add_break(WD_BREAK.PAGE)
