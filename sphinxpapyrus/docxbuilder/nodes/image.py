# -*- coding: utf-8 -*-
"""
Translate docutils node image formatting.
each image start will processed with visit() and finished with depart()
"""
import os

from docutils import nodes
from docutils.nodes import Element
from docutils.nodes import Node
from docx.enum.table import WD_TABLE_ALIGNMENT

from sphinxpapyrus.docxbuilder.translator import DocxTranslator

node_name = "image"


def visit(visitor: DocxTranslator, node: Element):
    """Start processing image node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Element)

    atts = {}
    uri = node['uri']
    ext = os.path.splitext(uri)[1].lower()
    if 'width' in node:
        atts['width'] = node['width']
    if 'height' in node:
        atts['height'] = node['height']
    if 'scale' in node:
        pass
    image_fullpath = os.path.join(visitor.builder.srcdir, uri)
    block_width = visitor.docx._block_width
    if isinstance(node.parent, nodes.substitution_definition):
        pass
    else:
        if isinstance(node.parent, nodes.paragraph):
            pic = visitor.r.add_picture(image_fullpath)
        elif isinstance(node.parent, nodes.figure):
            pic = visitor.r.add_picture(image_fullpath)
        else:
            p = visitor._add_paragraph()
            r = p.add_run()
            pic = r.add_picture(image_fullpath)
            align = node.get('align')
            if not align:
                align = visitor.builder.config.docx_imagetable_align
            if align:
                if align == 'left':
                    p.alignment = WD_TABLE_ALIGNMENT.LEFT
                elif align == 'center':
                    p.alignment = WD_TABLE_ALIGNMENT.CENTER
                elif align == 'right':
                    p.alignment = WD_TABLE_ALIGNMENT.RIGHT
        if pic.width > block_width:
            pic.height = int(pic.height * float(block_width) / pic.width)
            pic.width = block_width
    raise nodes.SkipNode


def depart(visitor: DocxTranslator, node: Node):
    """Finish processing image node"""
    assert isinstance(visitor, DocxTranslator)
    assert isinstance(node, Node)
