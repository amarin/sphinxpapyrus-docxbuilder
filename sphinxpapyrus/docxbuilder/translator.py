# -*- coding: utf-8 -*-
import glob
import os
from functools import partial
from importlib import import_module

from docutils import nodes
from docutils.nodes import Node
from docutils.nodes import NodeVisitor
from docutils.nodes import field_list
from docutils.nodes import node_class_names
from docutils.nodes import option_list
from docutils.nodes import table
from docx.document import Document
from sphinx.util import logging
from sphinx.util.console import bold
from sphinx.util.console import darkgreen

logger = logging.getLogger(__name__)

base_path = os.path.dirname(os.path.realpath(__file__))
directives_path = os.path.join(base_path, 'nodes')


class _DocxTranslator(NodeVisitor):
    """Implement translator methods"""

    def __init__(self, document, builder, docx):
        # import DocxBuilder locally to prevent circular imports
        self.tables = []
        from sphinxpapyrus.docxbuilder.builder import DocxBuilder
        # ensure input types
        assert isinstance(document, Node)
        assert isinstance(builder, DocxBuilder)
        assert isinstance(docx, Document)

        super().__init__(document)

        self._docx = docx
        self._builder = builder

        self._p = None  # current paragraph
        self._p_parents = [self.docx]  # current paragraph parents
        self._p_style = []  # current paragraph style('s)
        self._p_level = 0  # current paragraph level

        self._r = None  # current run
        self._r_style = None  # current run style

        self.settings = document.settings
        self.docnames = [builder.current_doc_name]

        self.numbered = 0
        self.numbered_level = 0
        self.section_level = 0
        self.section_numIds = [self._get_new_num(abstractNumId=12)]
        self.initial_header_level = 0

        self.numIds = []
        self.is_first_list_item = False
        # special paragraphs
        self.item_width_rate = 0.8

    def get_builder(self):
        """Take current builder instance"""
        return self._builder

    def get_p_parents(self):
        """Take current paragraph parents list"""
        return self._p_parents

    def get_docx(self) -> Document:
        """Take current docx"""
        return self._docx

    docx = property(get_docx)

    def get_r(self):
        """Take current run"""
        return self._r

    def set_r(self, new_r) -> None:
        """Set current run"""
        self._r = new_r

    def get_r_style(self):
        """Take current r_style"""
        return self._r_style

    def set_r_style(self, r_style) -> None:
        """Set current r_style"""
        self._r_style = r_style

    def get_p_style(self):
        """Take current p_style"""
        return self._p_style

    def set_p_style(self, p_style) -> None:
        """Set current p_style"""
        self._p_style = p_style

    def get_p_level(self) -> int:
        """Take current p_level"""
        return self._p_level

    def set_p_level(self, new_level: int) -> None:
        """Set current p_level"""
        self._p_level = new_level

    def get_p(self):
        """Take current para"""
        return self._p

    def set_p(self, new_p) -> None:
        """Set current para"""
        self._p = new_p

    def _fignum_prefix(self, node):
        prefix = ''
        if self.builder.config.numfig:
            figtype = self.builder.env.domains['std'].get_figtype(node)
            format = self.builder.config.numfig_format.get(figtype)
            alias = u'%s/%s' % (self.docnames[-1], figtype)
            id = node['ids'][0]
            nums = self.builder.fig_numbers[alias][id]
            prefix = format % '.'.join(map(str, nums))
        return prefix

    def _add_paragraph(self, text=None, style=None):
        p = None
        try:
            if isinstance(style, list):
                p = self.p_parents[-1].add_paragraph(text, style[-1])
            else:
                p = self.p_parents[-1].add_paragraph(text, style)
        except:
            p = self.p_parents[-1].add_paragraph(text, 'Normal')
        if self.p_level > 0:
            self._multilevel_list_numbering(p, self.p_level - 1, 15)
        return p

    def _add_run(self, text=None, style=None):
        r = None
        if self.p:
            try:
                r = self.p.add_run(text, style)
            except:
                r = self.p.add_run(text, 'Default Paragraph Font')
        return r

    def _get_new_num(self, abstractNumId):
        # monkey patch
        from types import MethodType
        from docx.oxml.numbering import CT_Num

        def add_num(self, abstractNum_id, restart=False):
            next_num_id = self._next_numId
            num = CT_Num.new(next_num_id, abstractNum_id)
            if restart:
                num.add_lvlOverride(ilvl=0).add_startOverride(1)
            return self._insert_num(num)

        if not self.docx:
            return 1

        assert isinstance(self._docx, Document)
        numbering = self.docx._part.numbering_part.numbering_definitions \
            ._numbering

        numbering.add_num = MethodType(add_num, numbering)

        num = numbering.add_num(abstractNumId, True).numId
        return num

    def _multilevel_list_numbering(self, paragraph, ilvl, numId):
        # monkey patch
        pfmt = paragraph.paragraph_format
        numPr = pfmt._element.get_or_add_pPr().get_or_add_numPr()
        numPr.get_or_add_ilvl().val = ilvl
        numPr.get_or_add_numId().val = numId

    def unknown_visit(self, node):
        # type: (nodes.Node) -> None
        raise NotImplementedError('Unknown node: ' + node.__class__.__name__)

    def _add_paragraph_between_table(self, node):
        index = node.parent.index(node)
        prev_node = node.parent[index - 1]

        if isinstance(prev_node, (table, field_list, option_list)):
            self.docx.add_paragraph()

    p = property(get_p, set_p)
    p_level = property(get_p_level, set_p_level)
    r = property(get_r, set_r)
    p_style = property(get_p_style, set_p_style)
    builder = property(get_builder)
    r_style = property(get_r_style, set_r_style)
    p_parents = property(get_p_parents)


class DocxTranslator(_DocxTranslator):

    def __init__(self, document, builder, docx):
        logger.info(bold(f"load translation functions ..."), nonl=True)
        # noinspection PyUnresolvedReferences
        for _path in glob.glob(os.path.join(directives_path, '*.py')):
            node_module_name = _path[len(base_path):-3]

            if node_module_name.endswith('__init__'):
                continue

            node_module_name = node_module_name.replace('/', '.')
            if node_module_name.startswith('.'):
                node_module_name = node_module_name[1:]

            node_module_name_parts = __name__.split('.')
            node_module_name_parts[-1] = node_module_name
            node_module_name = '.'.join(node_module_name_parts)

            node_name = node_module_name.split('.')[-1]

            try:
                imported_module = import_module(node_module_name)

                node_name = getattr(imported_module, 'node_name')
                visit_callable = getattr(imported_module, 'visit')
                depart_callable = getattr(imported_module, 'depart')

                visit_name = f'visit_{node_name}'
                depart_name = f'depart_{node_name}'

                # attach function as instance method using recipe from
                # https://stackoverflow.com/questions/30294458/
                # any-elegant-way-to-add-a-method-to-an-existing
                # -object-in-python
                method_visit = partial(visit_callable, self)
                method_visit.__name__ = visit_name
                setattr(self, visit_name, method_visit)

                method_depart = partial(depart_callable, self)
                method_depart.__name__ = depart_name
                setattr(self, depart_name, method_depart)

                logger.info(darkgreen(f" {node_name}"), nonl=True)

            except Exception as exc:
                logger.error(f"Cant import {node_name} translators: {exc}")
                raise exc
        logger.info(f" done")
        super().__init__(document, builder, docx)


translator_module_template = """# -*- coding: utf-8 -*-
'''
Translate docutils node %(node_name)s formatting.
each %(node_name)s start will processed with visit() and finished with depart()
'''

from docutils.nodes import Node
from %(visitor_package)s import %(visitor_class_name)s

node_name = "%(node_name)s"


def visit(visitor: %(visitor_class_name)s, node: Node):
    '''Start processing %(node_name)s node'''
    assert isinstance(visitor, %(visitor_class_name)s)
    assert isinstance(node, Node)


def depart(visitor: %(visitor_class_name)s, node: Node):
    '''Finish processing %(node_name)s node'''
    assert isinstance(visitor, %(visitor_class_name)s)
    assert isinstance(node, Node)
"""

if __name__ == '__main__':

    visitor_class = DocxTranslator
    visitor_package = 'sphinxpapyrus.docxbuilder.translator'

    for node in node_class_names:
        # noinspection PyUnresolvedReferences
        node_file_name = os.path.join(directives_path, f'{node.lower()}.py')

        if os.path.exists(node_file_name):
            continue

        with open(node_file_name, 'w') as node_module:
            node_module_code = translator_module_template % dict(
                node_name=node,
                visitor_class_name=visitor_class.__name__,
                visitor_package=visitor_package
            )
            node_module_code = node_module_code.replace("'''", '"""')
            node_module.write(node_module_code)
